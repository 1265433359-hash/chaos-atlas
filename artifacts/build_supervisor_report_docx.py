#!/usr/bin/env python3
"""Convert report_for_supervisor.md to a formatted Word document (A4, CJK-friendly).

Handles: #/##/### headings, markdown tables, fenced code blocks, "- " bullets,
"**bold**" inline, blockquotes "> ".
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\APP\project\chaos\artifacts\report_for_supervisor.md"
OUT = r"C:\Users\xiao junyang\Desktop\chaos-methodology-supervisor-report-v2.docx"

HEADING_COLOR = RGBColor(0x1F, 0x3B, 0x73)  # deep blue


def set_cjk(run, font_ascii="Calibri", font_cjk="微软雅黑", size=None):
    run.font.name = font_ascii
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_cjk)
    if size:
        run.font.size = Pt(size)


def add_runs(paragraph, text, base_size=None, bold_all=False):
    """Add text with **bold** markers to paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if bold_all:
                run.bold = True
        set_cjk(run, size=base_size)
    return paragraph


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def build():
    lines = open(SRC, encoding="utf-8").read().splitlines()
    doc = Document()

    # page setup A4 + margins
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)

    # base style
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    i = 0
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        for cl in code_buf:
            p = doc.add_paragraph()
            run = p.add_run(cl if cl else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            # light gray shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.space_after = Pt(0)
        code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_code()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            flush_code()
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_heading(level=level)
            p.clear()
            add_runs(p, text, base_size={1: 18, 2: 14, 3: 12}[level])
            for run in p.runs:
                run.font.color.rgb = HEADING_COLOR
                run.bold = True
            p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            flush_code()
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            j = i + 2
            rows = []
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            ncol = len(header_cells)
            table = doc.add_table(rows=1 + len(rows), cols=ncol)
            table.style = "Table Grid"
            for ci, h in enumerate(header_cells):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, h, base_size=9)
                for run in p.runs:
                    run.bold = True
                shade_cell(cell, "D5E8F0")
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    val = row[ci] if ci < len(row) else ""
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs(p, val, base_size=9)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i = j
            continue

        # bullet
        if re.match(r"^\s*[-*]\s+", line):
            flush_code()
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text, base_size=10.5)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            flush_code()
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            add_runs(p, text, base_size=10.5)
            for run in p.runs:
                run.italic = True
            p.paragraph_format.left_indent = Cm(0.6)
            i += 1
            continue

        # empty
        if not stripped:
            flush_code()
            i += 1
            continue

        # normal paragraph
        flush_code()
        p = doc.add_paragraph()
        add_runs(p, line, base_size=10.5)
        i += 1

    flush_code()
    doc.save(OUT)
    print("WROTE", OUT)


if __name__ == "__main__":
    build()
